from google import genai
import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import re
import os
from dotenv import load_dotenv
import PyPDF2
import pdfplumber
from collections import Counter
import math

load_dotenv()  
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file"""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def extract_keywords(text):
    """Extract important keywords from text"""
    # Convert to lowercase and remove special characters
    text = re.sub(r'[^\w\s]', ' ', text.lower())
    
    # Common stop words to exclude
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
        'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after',
        'above', 'below', 'between', 'among', 'is', 'are', 'was', 'were', 'be', 'been',
        'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
        'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those',
        'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your',
        'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', 'her',
        'hers', 'herself', 'it', 'its', 'itself', 'they', 'them', 'their', 'theirs',
        'themselves', 'what', 'which', 'who', 'whom', 'whose', 'where', 'when', 'why',
        'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some',
        'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very'
    }
    
    # Split into words and filter
    words = text.split()
    keywords = [word for word in words if len(word) > 2 and word not in stop_words]
    
    # Also extract multi-word phrases (2-3 words)
    phrases = []
    for i in range(len(words) - 1):
        if len(words[i]) > 2 and len(words[i+1]) > 2:
            phrase = f"{words[i]} {words[i+1]}"
            if not any(stop in phrase for stop in ['the ', ' the', 'and ', ' and']):
                phrases.append(phrase)
    
    # Combine single words and phrases
    all_keywords = keywords + phrases
    
    # Count frequency and return most common
    keyword_counts = Counter(all_keywords)
    return keyword_counts

def calculate_ats_score(resume_text, job_description):
    """Calculate ATS compatibility score between resume and job description"""
    
    # Extract keywords from both texts
    resume_keywords = extract_keywords(resume_text)
    job_keywords = extract_keywords(job_description)
    
    # Get top keywords from job description (more weight to frequent terms)
    top_job_keywords = dict(job_keywords.most_common(50))
    
    # Calculate matches
    matched_keywords = []
    missed_keywords = []
    total_job_keyword_weight = sum(top_job_keywords.values())
    matched_weight = 0
    
    for keyword, job_freq in top_job_keywords.items():
        if keyword in resume_keywords:
            matched_keywords.append({
                'keyword': keyword,
                'job_freq': job_freq,
                'resume_freq': resume_keywords[keyword]
            })
            matched_weight += job_freq
        else:
            missed_keywords.append({
                'keyword': keyword,
                'job_freq': job_freq
            })
    
    # Calculate base score (percentage of matched keyword weight)
    if total_job_keyword_weight > 0:
        base_score = (matched_weight / total_job_keyword_weight) * 100
    else:
        base_score = 0
    
    # Bonus points for exact phrase matches
    job_text_lower = job_description.lower()
    resume_text_lower = resume_text.lower()
    
    # Look for exact skill matches
    tech_skills = [
        'python', 'javascript', 'react', 'node.js', 'sql', 'html', 'css',
        'machine learning', 'data analysis', 'project management', 'agile',
        'scrum', 'git', 'docker', 'kubernetes', 'aws', 'azure', 'gcp'
    ]
    
    skill_bonus = 0
    for skill in tech_skills:
        if skill in job_text_lower and skill in resume_text_lower:
            skill_bonus += 2
    
    # Calculate final score (cap at 100)
    final_score = min(base_score + skill_bonus, 100)
    
    return {
        'score': round(final_score, 1),
        'matched_keywords': matched_keywords[:10],  # Top 10 matches
        'missed_keywords': missed_keywords[:10],    # Top 10 missed
        'total_job_keywords': len(top_job_keywords),
        'total_matched': len(matched_keywords),
        'suggestions': generate_ats_suggestions(final_score, missed_keywords[:5])
    }

def generate_ats_suggestions(score, missed_keywords):
    """Generate suggestions based on ATS score"""
    suggestions = []
    
    if score < 50:
        suggestions.append("🔴 Low ATS Score: Your resume needs significant optimization")
        suggestions.append("📝 Add more relevant keywords from the job description")
        suggestions.append("🎯 Focus on matching the required skills and experience")
    elif score < 70:
        suggestions.append("🟡 Moderate ATS Score: Good foundation, but room for improvement")
        suggestions.append("✨ Fine-tune your resume to include more specific terms")
    else:
        suggestions.append("🟢 Good ATS Score: Your resume is well-optimized!")
        suggestions.append("🚀 Minor tweaks could make it even better")
    
    if missed_keywords:
        top_missed = [kw['keyword'] for kw in missed_keywords[:3]]
        suggestions.append(f"🎯 Try to include these important terms: {', '.join(top_missed)}")
    
    return suggestions

def parse_markdown_line(line):
    """Parse a line and return formatted parts"""
    parts = []
    
    # Handle nested formatting like ***Languages:** or **Languages:**
    # Pattern: one or more * followed by text, then : (optional), then more *
    pattern = r'(\*{2,3})([^*:]+):?(\*{2,3})?'
    
    current_pos = 0
    for match in re.finditer(pattern, line):
        # Add text before the match as regular text
        if match.start() > current_pos:
            parts.append({
                'text': line[current_pos:match.start()],
                'type': 'regular'
            })
        
        # Add the formatted text
        stars_before = match.group(1)
        text = match.group(2).strip()
        stars_after = match.group(3) if match.group(3) else ''
        
        # Determine if it's a subheading (usually **text:** or ***text:**)
        if len(stars_before) >= 2:
            parts.append({
                'text': text + ':' if ':' in match.group(0) else text,
                'type': 'subheading'
            })
        else:
            parts.append({
                'text': text,
                'type': 'bold'
            })
        
        current_pos = match.end()
    
    # Add remaining text
    if current_pos < len(line):
        remaining_text = line[current_pos:].strip()
        if remaining_text:
            parts.append({
                'text': remaining_text,
                'type': 'regular'
            })
    
    # If no patterns found, treat as regular text
    if not parts:
        parts.append({
            'text': line,
            'type': 'regular'
        })
    
    return parts

def generate_custom_resume(resume_text, job_description):
    prompt = f"""
    You are a professional resume editor. Customize the resume below to better fit the job description.

    Resume:
    {resume_text}

    Job Description:
    {job_description}

    Return a new version of the resume tailored to the job.
    """

    response = client.models.generate_content(
        model="gemini-2.5-flash", contents=prompt
    )
    return response.text

def convert_text_to_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(3)  # Add space for empty lines
            continue
            
        # Check if line is a main heading (surrounded by **)
        if line.startswith('**') and line.endswith('**') and '***' not in line:
            # It's a main heading
            heading_text = line[2:-2].strip()
            pdf.set_font("Arial", 'B', size=14)
            try:
                heading_encoded = heading_text.encode('latin-1', 'ignore').decode('latin-1')
                pdf.multi_cell(0, 12, heading_encoded)
            except:
                heading_clean = ''.join(char for char in heading_text if ord(char) < 128)
                pdf.multi_cell(0, 12, heading_clean)
            pdf.ln(2)
        else:
            # Parse the line for mixed formatting
            parts = parse_markdown_line(line)
            current_line = ""
            
            for part in parts:
                if part['type'] == 'subheading':
                    # If we have accumulated text, output it first
                    if current_line:
                        pdf.set_font("Arial", size=11)
                        try:
                            current_line_encoded = current_line.encode('latin-1', 'ignore').decode('latin-1')
                            pdf.multi_cell(0, 7, current_line_encoded)
                        except:
                            current_line_clean = ''.join(char for char in current_line if ord(char) < 128)
                            pdf.multi_cell(0, 7, current_line_clean)
                        current_line = ""
                    
                    # Output subheading
                    pdf.set_font("Arial", 'B', size=12)
                    try:
                        subheading_encoded = part['text'].encode('latin-1', 'ignore').decode('latin-1')
                        pdf.multi_cell(0, 8, subheading_encoded)
                    except:
                        subheading_clean = ''.join(char for char in part['text'] if ord(char) < 128)
                        pdf.multi_cell(0, 8, subheading_clean)
                else:
                    # Accumulate regular and bold text
                    current_line += part['text']
            
            # Output any remaining text
            if current_line:
                pdf.set_font("Arial", size=11)
                try:
                    current_line_encoded = current_line.encode('latin-1', 'ignore').decode('latin-1')
                    pdf.multi_cell(0, 7, current_line_encoded)
                except:
                    current_line_clean = ''.join(char for char in current_line if ord(char) < 128)
                    pdf.multi_cell(0, 7, current_line_clean)

    return pdf.output(dest='S').encode('latin-1')

def convert_text_to_docx(text):
    doc = Document()
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()  # Add empty paragraph for spacing
            continue
            
        # Check if line is a main heading (surrounded by ** but not ***)
        if line.startswith('**') and line.endswith('**') and '***' not in line:
            # It's a main heading
            heading_text = line[2:-2].strip()
            heading = doc.add_heading(heading_text, level=2)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            # Parse the line for mixed formatting
            parts = parse_markdown_line(line)
            paragraph = doc.add_paragraph()
            
            for part in parts:
                run = paragraph.add_run(part['text'])
                if part['type'] == 'subheading':
                    run.bold = True
                elif part['type'] == 'bold':
                    run.bold = True
    
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.set_page_config(page_title="AI Resume Customizer", layout="centered")
st.title("🧠 AI Resume Customizer")
st.markdown("Upload your base resume and paste the job description. Edit the AI-generated result and download it as PDF or Word!")

# Upload and read resume file
resume_text = ""
uploaded_file = st.file_uploader("Upload your base resume", type=["txt", "pdf"])
if uploaded_file is not None:
    if uploaded_file.type == "text/plain":
        resume_text = uploaded_file.read().decode("utf-8")
    elif uploaded_file.type == "application/pdf":
        resume_text = extract_text_from_pdf(uploaded_file)
        if resume_text is None:
            st.error("Failed to extract text from PDF. Please try uploading a text file instead.")
    else:
        st.error("Unsupported file type. Please upload a .txt or .pdf file.")

# Input: Job Description
job_description = st.text_area("Paste Job Description", height=200)

if st.button("Generate Tailored Resume") and resume_text and job_description:
    with st.spinner("Generating resume using Gemini..."):
        new_resume = generate_custom_resume(resume_text, job_description)
        st.success("Resume generated successfully!")

        # Store the generated resume in session state
        st.session_state.generated_resume = new_resume

# Show edit area and download buttons if resume is generated
if 'generated_resume' in st.session_state:
    
    # Add ATS Score Analysis
    st.subheader("📊 ATS Compatibility Analysis")
    
    if st.button("🔍 Analyze ATS Score"):
        with st.spinner("Analyzing ATS compatibility..."):
            ats_results = calculate_ats_score(st.session_state.generated_resume, job_description)
            
            # Display score with color coding
            score = ats_results['score']
            if score >= 70:
                st.success(f"🎯 ATS Score: {score}%")
            elif score >= 50:
                st.warning(f"⚠️ ATS Score: {score}%")
            else:
                st.error(f"❌ ATS Score: {score}%")
            
            # Create columns for detailed analysis
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**✅ Matched Keywords:**")
                for match in ats_results['matched_keywords']:
                    st.write(f"• {match['keyword']} (appears {match['resume_freq']} times)")
            
            with col2:
                st.write("**❌ Missing Keywords:**")
                for missed in ats_results['missed_keywords']:
                    st.write(f"• {missed['keyword']} (important in job desc)")
            
            # Show suggestions
            st.write("**💡 Suggestions:**")
            for suggestion in ats_results['suggestions']:
                st.write(suggestion)
            
            # Store ATS results in session state for re-display
            st.session_state.ats_results = ats_results
    
    # Display previous ATS results if available
    if 'ats_results' in st.session_state:
        with st.expander("📈 Previous ATS Analysis Results"):
            results = st.session_state.ats_results
            score = results['score']
            
            # Score with progress bar
            st.metric("ATS Compatibility Score", f"{score}%")
            st.progress(score / 100)
            
            col1, col2, col3 = st.columns(3)
            col1.metric("Total Job Keywords", results['total_job_keywords'])
            col2.metric("Keywords Matched", results['total_matched'])
            col3.metric("Match Rate", f"{(results['total_matched']/results['total_job_keywords']*100):.1f}%" if results['total_job_keywords'] > 0 else "0%")
    
    st.divider()
    
    # Editable field
    edited_resume = st.text_area("✏️ Edit Your Resume", value=st.session_state.generated_resume, height=400)
    
    # Create two columns for download buttons
    col1, col2 = st.columns(2)
    
    with col1:
        # PDF Download
        try:
            pdf_bytes = convert_text_to_pdf(edited_resume)
            st.download_button(
                label="📄 Download as PDF",
                data=pdf_bytes,
                file_name="tailored_resume.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.error(f"Error generating PDF: {str(e)}")
    
    with col2:
        # DOCX Download
        try:
            docx_buffer = convert_text_to_docx(edited_resume)
            st.download_button(
                label="📝 Download as DOCX",
                data=docx_buffer.getvalue(),
                file_name="tailored_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Error generating DOCX: {str(e)}")

elif not resume_text:
    st.info("Please upload a resume file.")
elif not job_description:
    st.info("Please paste a job description.")